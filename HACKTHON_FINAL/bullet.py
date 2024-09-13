from docx import Document
import pyautogui
import pyperclip,time
from doc import fname
import stack

# Load the existing .docx file or create a new one


# Function to append bulleted text
def append_bulleted_line(text):
    
    file_path = fname()
    try:
        doc = Document(file_path)
    except:
        doc = Document()
    # Add a paragraph with bullet style
    sstyle="ListBullet"
    paragraph=doc.add_paragraph(text, style=sstyle)
    # Save the updated document
    stack.set_action(paragraph)
    doc.save(file_path) 
    # return paragraph

def add_bullet_to_doc(parent):
    global root
    parent.destroy()
    root = None
    # Copy the selected text to the clipboard (assumes the user has selected text)
    pyautogui.hotkey('ctrl', 'c')
    time.sleep(0.1)  # Small delay to ensure the clipboard is updated
    selected_text = pyperclip.paste()
    
    if selected_text:
        append_bulleted_line(selected_text)

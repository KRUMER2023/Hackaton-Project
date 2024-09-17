from pathlib import Path
from tkinter import Tk, Canvas, Entry, Text, Button, PhotoImage
from pynput import mouse
import tkinter as tk
import threading
import pystray
import PIL.Image
import pyautogui
import pyperclip
import time  # Import the time module
from docx import Document
from push import create_push
from h1 import add_h1_to_doc
from h2 import add_h2_to_doc
from add import add_text_to_doc
from bullet import add_bullet_to_doc
from Ai_gui import ai_menu
from doc import fname
from more import more_gui2
import stack



OUTPUT_PATH = Path(__file__).parent
ASSETS_PATH = OUTPUT_PATH / Path(r"assets2\frame0")

def relative_to_assets(path: str) -> Path:
    return ASSETS_PATH / Path(path)

dragging = False #dragging ko check kar ta hai
window =None

# Function to append bulleted text
def append_bulleted_line(text):
    file_path = fname
    doc = Document(file_path)
    paragraph = doc.add_paragraph(text, style='ListBullet')
    stack.set_action(paragraph)
    return paragraph

# Function to undo the last action (removes the last appended text)
def undo_last_action():
    action_stack=stack.get_action()
    if action_stack:
        last_action = stack.pop_action()
        stack.set_redo(last_action)  # Keep track of undone actions for redo
        last_action._element.getparent().remove(last_action._element)  # Remove the paragraph
        print("Undo successful.")
    else:
        print("No action to undo.")

# Function to redo the last undone action
def redo_last_action():
    file_path = fname
    doc = Document(file_path)
    redo_stack=stack.get_redo()
    if redo_stack:
        last_redo = stack.pop_redu()
        stack.set_action(last_redo)
        doc._body.append(last_redo._element)  # Reinsert the paragraph
        print("Redo successful.")
    else:
        print("No action to redo.")


def show_popup():
    
    # Get the current cursor position
    x, y = pyautogui.position()    
    global window # Reference the global add_icon variable
    
    # Destroy the previous "Push" popup if it exists
    if window is not None:
        try:
            window.destroy()
        except tk.TclError:
            pass  # Ignore errors if the window is already destroyed
        
        
    window = Tk()
    window.overrideredirect(True)  # Remove window decorations (title bar, etc.)
    window.attributes('-topmost', True)  # Keep the window on top
    
    # Set the window position (right side of the cursor)
    window.geometry(f"+{x + 10}+{y}")

    window.geometry("381x55")
    window.configure(bg = "#070707")


    canvas = Canvas(
        window,
        bg = "#070707",
        height = 55,
        width = 381,
        bd = 0,
        highlightthickness = 0,
        relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    image_image_1 = PhotoImage(
        file=relative_to_assets("image_1.png"))
    image_1 = canvas.create_image(
        190.0,
        27.0,
        image=image_image_1
    )

    # 
    button_image_1 = PhotoImage(
        file=relative_to_assets("button_1.png"))
    button_1 = Button(
        image=button_image_1,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: more_gui2(window),
        relief="flat"
    )
    button_1.place(
        x=323.0,
        y=11.0,
        width=41.0,
        height=36.0
    )

    button_image_2 = PhotoImage(
        file=relative_to_assets("button_2.png"))
    button_2 = Button(
        image=button_image_2,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: redo_last_action(),
        relief="flat"
    )
    button_2.place(
        x=279.0,
        y=10.0,
        width=41.0,
        height=37.0
    )

    button_image_3 = PhotoImage(
        file=relative_to_assets("button_3.png"))
    button_3 = Button(
        image=button_image_3,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: undo_last_action(),
        relief="flat"
    )
    button_3.place(
        x=235.0,
        y=10.0,
        width=41.0,
        height=37.0
    )

    button_image_4 = PhotoImage(
        file=relative_to_assets("button_4.png"))
    button_4 = Button(
        image=button_image_4,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: ai_menu(window),
        relief="flat"
    )
    button_4.place(
        x=192.0,
        y=11.0,
        width=39.0,
        height=34.0
    )

    button_image_5 = PhotoImage(
        file=relative_to_assets("button_5.png"))
    button_5 = Button(
        image=button_image_5,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: add_bullet_to_doc(window),
        relief="flat"
    )
    button_5.place(
        x=148.0,
        y=11.0,
        width=39.0,
        height=34.0
    )

    button_image_6 = PhotoImage(
        file=relative_to_assets("button_6.png"))
    button_6 = Button(
        image=button_image_6,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: add_h2_to_doc(window),
        relief="flat"
    )
    button_6.place(
        x=104.0,
        y=11.0,
        width=39.0,
        height=34.0
    )

    button_image_7 = PhotoImage(
        file=relative_to_assets("button_7.png"))
    button_7 = Button(
        image=button_image_7,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: add_h1_to_doc(window),
        relief="flat"
    )
    button_7.place(
        x=60.0,
        y=11.0,
        width=39.0,
        height=34.0
    )

    # SIMPLE TEXT BUTTON
    button_image_8 = PhotoImage(
        file=relative_to_assets("button_8.png"))
    button_8 = Button(
        image=button_image_8,
        borderwidth=0,
        highlightthickness=0,
        command=lambda:add_text_to_doc(window),
        relief="flat"
    )
    button_8.place(
        x=16.0,
        y=11.0,
        width=39.0,
        height=34.0
    )
    window.resizable(False, False)
    window.mainloop()

def on_click(x, y, button, pressed):
    global dragging,window
    if button == mouse.Button.left:
        if pressed:
            dragging = False  # Reset dragging state on mouse press
        else:
            if dragging:  # Only show popup if dragging occurred
                # show_popup(folder,doc_name)
                show_popup()

def on_move(x, y):
    global dragging
    dragging = True  # If the mouse is moved while pressed, set dragging to True

def popup_main():
    
    # icon=pystray.Icon("Temp",image,menu=pystray.Menu(
    #     pystray.MenuItem("Continue",on_tray_clicked),
    #     pystray.MenuItem("Pause",on_tray_clicked),
    #     pystray.MenuItem("Update",on_tray_clicked),
    #     pystray.MenuItem("Exit",on_tray_clicked)
    # ))

    
    # Start listening to mouse events
    with mouse.Listener(on_click=on_click, on_move=on_move) as listener:
            listener.join()
            
# popup_main()

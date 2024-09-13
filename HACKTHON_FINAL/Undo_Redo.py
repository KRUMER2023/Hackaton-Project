from docx import Document

# Load the existing .docx file or create a new one
file_path = 'your_file.docx'
try:
    doc = Document(file_path)
except:
    doc = Document()


global action_stack,redo_stack
# To track the changes
action_stack = []
redo_stack = []

# Function to append bulleted text
def append_bulleted_line(text):
    paragraph = doc.add_paragraph(text, style='ListBullet')
    action_stack.append(paragraph)
    return paragraph

# Function to undo the last action (removes the last appended text)
def undo_last_action():
    if action_stack:
        last_action = action_stack.pop()
        redo_stack.append(last_action)  # Keep track of undone actions for redo
        last_action._element.getparent().remove(last_action._element)  # Remove the paragraph
        print("Undo successful.")
    else:
        print("No action to undo.")

# Function to redo the last undone action
def redo_last_action():
    if redo_stack:
        last_redo = redo_stack.pop()
        action_stack.append(last_redo)
        doc._body.append(last_redo._element)  # Reinsert the paragraph
        print("Redo successful.")
    else:
        print("No action to redo.")

# Main function to control user input
def UR():
    while True:
        print("\nOptions: 1) Append new line  2) Undo  3) Redo  4) Save & Exit")
        choice = input("Choose an option (1/2/3/4): ").strip()

        if choice == "1":
            text_to_append = input("Enter the text to append: ").strip()
            append_bulleted_line(text_to_append)
            redo_stack.clear()  # Clear redo history after new append
            print(f"Text appended successfully.")

        elif choice == "2":
            undo_last_action()

        elif choice == "3":
            redo_last_action()

        elif choice == "4":
            doc.save(file_path)
            print(f"Document saved successfully to {file_path}")
            break

        else:
            print("Invalid option, please choose again.")
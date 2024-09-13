import tkinter as tk
import pyautogui
from docx import Document
from docx.shared import Pt
import time  # Import the time module
import pyperclip
from doc import file_maker,fname
import stack


button_add=tk
dragging = False


def append_text_to_doc(text):
    name= fname()
    doc = Document(name)
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(14)  # Set font size to 48 points (96 half-points), similar to H1
    doc.save(name)
    stack.set_action(text)
    print(f"Appended text: {text}")
    

def add_h1_to_doc(parent):
    global root
    parent.destroy()
    root = None
    # Copy the selected text to the clipboard (assumes the user has selected text)
    pyautogui.hotkey('ctrl', 'c')
    time.sleep(0.1)  # Small delay to ensure the clipboard is updated
    selected_text = pyperclip.paste()
    
    if selected_text:
        append_text_to_doc(selected_text)

# def create_h1(parent):
#     # Create a button with the text "Push"
#     button_push = tk.Button(parent, text="H1", bg="black", fg="white", font=("Arial", 7), padx=10, pady=5, command=lambda: add_text_to_doc(parent))
#     button_push.pack(side=tk.LEFT)
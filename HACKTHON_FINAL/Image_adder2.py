import tkinter as tk
import pyautogui
import pyperclip
from doc import file_maker,fname
from docx import Document
from docx.shared import Inches
import doc
import pyautogui
import time
from time import strftime
import datetime,os
from docx import Document
import os
from PIL import ImageGrab, Image
from io import BytesIO

def appender(name,strTime):
    doc = Document(name)
    # doc.add_picture(strTime+".png")
    picture_path = "IMAGES/"+strTime + ".png"
    # Add the picture to the document
    inline_shape= doc.add_picture(picture_path)

        # Resize the image to fit within the page
    width, height = inline_shape.width, inline_shape.height
    max_width, max_height = Inches(6), Inches(4)  # adjust these values to fit your page size
    if width > max_width or height > max_height:
        scale_factor = min(max_width / width, max_height / height)
        inline_shape.width, inline_shape.height = int(width * scale_factor), int(height * scale_factor)
    
    doc.save(name)
    
    os.remove("IMAGES/"+strTime + ".png")

def image_add(parent):
    
    parent.destroy()
    time.sleep(0.2)
    
    name= fname()
    
    strTime = datetime.datetime.now().strftime("%H-%M-%S")
    
    # Clear the clipboard
    pyperclip.copy('')

    # Simulate print screen button press
    pyautogui.press('printscreen')
    
    # Wait for the user to take a snip and save it to the clipboard
    while True:
        img = ImageGrab.grabclipboard()
        if img is not None:
            break
        time.sleep(0.1)
        
    time.sleep(1)
    try:
        # Grab the image from the clipboard
        image = ImageGrab.grabclipboard()

        if isinstance(image, Image.Image):
            # Save the image to the specified path
            image.save("IMAGES/"+strTime+'.png')
            # print(f"Image saved to {save_path}")
        else:
            print("No image found in the clipboard.")

    except Exception as e:
        print(f"Error: {e}")

    appender(name,strTime)
    
    
# def create_Photo(parent):
# # Create a button with the text "Push"
#     button_push = tk.Button(parent, text="◙", bg="black", fg="white", font=("Arial", 7), padx=10, pady=5, command=lambda: image_add(parent))
#     button_push.pack(side=tk.LEFT)
